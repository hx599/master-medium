# -*- coding: utf-8 -*-
"""精简表2/3/4的分析段落。"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document(r'D:\Master_medium\中期实验报告_修订版.docx')


def replace_para_containing(doc, search, new_text):
    for para in doc.paragraphs:
        if search in para.text:
            for run in para.runs:
                run.text = ''
            if para.runs:
                run = para.runs[0]
                run.text = new_text
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
            else:
                run = para.add_run(new_text)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
            return True
    return False


# ── 表2-4 总引语（精简）──
replace_para_containing(doc,
    '表2至表4展示了Houston、MUUFL和Berlin三个数据集上，选定未知类别组合的逐类分类精度',
    '表2至表4展示了三个数据集上选定未知类别组合的逐类精度。传统方法（ExViT、MFT）在所有组合中对未知类的精度均为零，因其不具备利用外部语义知识进行零样本推理的机制。UC方法在不同组合下的表现分化明显，以下选择若干典型组合进行分析。'
)

# ── Houston 表2（精简）──
replace_para_containing(doc,
    'UC在高速路上的精度为66.52%，在被压草上为73.82%',
    'Houston数据集上，"被压草"（第2类）与"高速路"（第10类）组合的H(aa)=76.79%，两类地物视觉差异显著，CLIP语义匹配较为准确。但"树木"（第4类）与"商业区"（第8类）组合的H(aa)仅43.15%，且高速公路作为已知类的精度降至8.20%，说明未知类伪标签中的噪声通过光谱修正环节波及了视觉相似的已知类。"铁路"（第11类）与"跑道"（第15类）组合的H(aa)=70.16%，虽然两类均为灰色线性结构，但空间上下文差异（铁路贯穿多个区域、跑道局限于固定场地）为区分提供了线索。ExViT和MFT在已知类上均保持90%以上精度（如(4,8)组合下ExViT对居民区达97.49%），其瓶颈完全在于缺乏零样本分类机制。'
)

# ── MUUFL 表3（精简）──
replace_para_containing(doc,
    'UC在树木上取得85.91%的精度，但水体精度为0%',
    'MUUFL数据集上，"树木"（第1类）与"水"（第6类）组合中水体精度为0%——水体在LiDAR近红外波段的强吸收导致回波缺失，与CLIP基于光学图像学到的水体表征严重失配。"以草地为主"（第2类）与"黄色路缘"（第10类）组合中黄路缘精度为0%，其像素占比极小，超像素分割难以独立提取。"裸土和沙地"（第4类）与"布制覆盖板"（第11类）组合中布制覆盖板同样为0%，该地物在CLIP训练数据中几乎不存在对应的图像-文本对。已知类方面，UC在建筑和公路上的精度分别为86.84%和82.49%，仍低于ExViT（98.01%）和MFT（97.07%）。'
)

# ── Berlin 表4（精简）──
replace_para_containing(doc,
    'UC在土壤上精度为66.72%，但家庭菜园精度为0%',
    'Berlin数据集上，"土壤"（第5类）与"家庭菜园"（第6类）组合的OA达91.59%，但家庭菜园精度为0%。"土壤"（第5类）与"水"（第8类）组合中，水体同样为0%。水体和家庭菜园在SAR图像中均缺乏稳定的散射特征——水体镜面反射的低回波易与平整裸地混淆，家庭菜园的体散射与森林等植被高度重叠。两者的共同症结在于CLIP缺乏对SAR成像机制的物理感知。"森林"（第1类）与"家庭菜园"（第6类）组合中UC的OA降至86.30%，森林精度仅36.93%。(5,6)组合下UC在居民区和工业区上的已知类精度均超90%，但未知类精度分别为66.72%和0%，性能落差完全源于CLIP伪标签在特定类别上的质量缺陷。'
)

# ── 总结（精简）──
replace_para_containing(doc,
    '从上述逐类分析可以看出，CLIP伪标签的质量受三类因素制约较为明显',
    '上述结果表明，CLIP伪标签质量主要受三类因素制约：一是地物空间尺度——面积小或分散的地物（黄色路缘）难以被超像素分割独立提取；二是辅助模态类型——SAR与CLIP光学先验之间的领域鸿沟显著大于LiDAR，Berlin数据集多个零精度类别即源于此；三是类别命名的语义覆盖——allotment（家庭菜园）、cloth panels（布制覆盖板）等罕见词汇无法触发CLIP的有效语义匹配。'
)

output_path = r'D:\Master_medium\中期实验报告_精简版.docx'
doc.save(output_path)
print(f'精简完成！已保存至 {output_path}')
