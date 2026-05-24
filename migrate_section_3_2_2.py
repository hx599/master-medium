# -*- coding: utf-8 -*-
"""把 §3 / §3.2 / §3.2.2 写入 中期实验报告_最终版.docx，
插在已有 §2.5 (...构建能够扩展类别空间并保持稳定性能的多模态遥感分类框架。) 之后、
原"表1显示了两个未知类别情况下的各方法平均分类结果"段落之前。

§3.2.2 严格按开题报告 §3.2.2 的连续行文风格（无加粗副标题）。
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SRC_DOCX = Path('/Users/agiuser/master-medium/中期实验报告_最终版.docx')
DST_DOCX = SRC_DOCX
ANCHOR_TEXT_PREFIX = '表1显示了两个未知类别情况下的各方法平均分类结果'


# (level, text)
# level: 'h1' | 'h2' | 'h3' | 'p' | 'fig' | 'cap' | 'eq'
SECTIONS = [
    ('h1', '3. 主要研究内容及研究方案'),
    ('p',
     '针对多模态遥感图像在实际分类应用中面临的对标注样本依赖性强、模态间特征差异显著以及难以有效识别未知类别等挑战，'
     '本课题深入研究视觉语言模型辅助下多模态遥感图像像素级分类，旨在提升模型的跨模态融合能力与未知类别识别能力，'
     '并增强其在动态环境中的持续适应性。'),

    ('h2', '3.2 研究方案'),
    ('p',
     '本课题的研究方案分为三个子方向：基于Transformer的多模态遥感图像预训练方法、'
     '基于视觉语言模型的多模态遥感图像未知类分类方法以及基于持续学习的多模态遥感图像未知类分类方法。'
     '本节作为中期工作的主体，重点说明第二个子方向的研究方案。'),

    ('h3', '3.2.2 基于视觉语言模型的多模态遥感图像未知类分类方法研究'),

    ('p',
     '在多模态遥感图像未知类分类任务中，核心挑战在于在缺乏未知类训练样本的情况下，'
     '如何获取未知类的先验知识并建立已知类与未知类之间的语义关联。'
     '现有多模态遥感数据面临标注不统一、类别规模有限、跨模态特征分布差异显著等问题，'
     '使得常规分类方法难以直接适用。为了实现未知类分类，需要解决两个关键问题：'
     '1）如何在已知类标签和多模态样本特征之间建立稳定的映射关系；'
     '2）如何在已知类与未知类之间建立语义层面的关联，从而获得未知类的特征表示与语义信息。'),

    ('p',
     '针对上述问题，本节提出基于视觉语言模型的多模态遥感图像未知类分类框架（MRS-VLM-Tr），如图3-3所示。'
     '该框架以视觉语言模型蕴含的语义先验为基础，通过伪标签获取机制将其语义知识引入像素级多模态分类任务，'
     '并结合监督训练使模型在保持已知类高精度分类的同时，具备对未知类的识别能力。'
     '整体上，框架由多模态伪标签生成与增强、谱特征引导的伪标签修正以及两阶段全局训练三个串联模块构成。'
     '设主模态高光谱图像为X_primary∈ℝ^(H×W×C₁)，辅助模态为X_aux∈ℝ^(H×W×C₂)，'
     '类别集合C=C_seen∪C_unseen由已知类和未知类组成；在训练阶段，C_unseen的样本标签全部被屏蔽，'
     '仅参与伪标签分支的监督。'),

    ('fig', '[图 3-3 占位]'),
    ('cap', '图3-3　基于视觉语言模型的多模态遥感图像未知类分类总体框架'),

    ('p',
     '在伪标签生成阶段，本节采用CLIP模型作为核心的视觉语言模型。'
     '对于一幅输入图像x与K个候选类别构成的文本描述集合{y_i^T}_{i=1}^K，'
     'CLIP通过对图像编码和文本编码计算余弦相似度，并经Softmax归一化输出该图像属于各类别的概率向量：'),

    ('eq', 'P = CLIP(x, {y_i^T}_{i=1}^K) ∈ ℝ^K', '(3-1)'),

    ('p',
     '由于CLIP在大规模自然图像-文本对上预训练得到，非光学辅助模态（LiDAR强度图、SAR散射图）'
     '在成像机理与统计特性上与之存在显著差异，直接调用难以获得可靠的语义匹配。'
     '因此本节将视觉语言模型的输入仅限于主模态高光谱图像，'
     '先按预设波段索引将X_primary降维为三通道RGB图像X_primary^RGB；'
     '同时，CLIP工作于图像级，而本任务为像素级，需要借助分割完成两者之间的过渡。'
     '本节在X_primary^RGB上联合采用ERS与SAM两种分割算法将整幅图像划分为N个语义连贯且像素全覆盖的区域：'),

    ('eq', '{s_i}_{i=1}^N = Seg_{ERS+SAM}(X_primary^RGB)', '(3-2)'),

    ('p',
     '考虑到单一视觉语言模型在遥感场景下可能存在语义偏置，'
     '本节拟引入两个具有互补语义先验的视觉语言模型，'
     '对每个区域s_i的代表图像块a_i分别推理后在概率层做平均融合：'),

    ('eq', 'P_i = (1/M) Σ_{m=1}^M CLIP_m(a_i, {y_j^T}_{j=1}^K)', '(3-3)'),

    ('p',
     '其中M=2。将P_i分配给区域s_i内的所有像素，最终形成整图层面的伪标签概率张量P∈ℝ^(H×W×K)。'),

    ('p',
     'CLIP生成的伪标签不可避免地存在噪声，其来源既包括视觉语言模型的语义先验与遥感地物分布之间的知识鸿沟，'
     '也包括超像素分割边界处的标签歧义。直接以噪声伪标签监督模型训练会污染特征学习。'
     '为此，本节引入谱特征引导的伪标签修正机制：'
     '先将像素级伪标签和多模态光谱特征聚合到超像素级，'
     '再对超像素特征进行谱聚类以捕捉光谱空间中的内在分布；'
     '对于每个超像素，通过比较其与同簇且同伪标签邻居的相似度权重之和'
     '与全图同伪标签邻居权重之和的比值得到局部一致性分数，'
     '当该分数低于预设阈值时，按P中的概率排序迭代替换为次高候选类，'
     '直至一致性恢复或候选类穷尽，从而得到修正后的伪标签P_c∈ℝ^(H×W×K)。'
     '最后，对包含已知监督标签的超像素直接以真值覆盖伪标签，'
     '确保已知类信号不被修正过程污染。'),

    ('p',
     '在全局训练阶段，本节采用两阶段策略对像素级分割网络进行优化。'
     '第一阶段仅以已知类真值标签Y对网络进行预训练，使其形成稳定的已知类特征表征；'
     '第二阶段以真值标签和修正后伪标签P_c联合监督网络，'
     '通过复合损失函数同时优化已知类的精确分类与未知类的语义识别能力：'),

    ('eq', 'L = L_sup + β · L_pl', '(3-4)'),

    ('p',
     '其中L_sup为基于真值标签计算的监督交叉熵损失，'
     'L_pl为基于修正后伪标签计算的伪标签交叉熵损失，β为伪标签损失权重。'
     '在网络输入端，本节采用早期融合策略，将X_primary与X_aux分别经过通道维标准化后'
     '沿通道维拼接为联合输入张量：'),

    ('eq', 'X_fused = Concat(Norm(X_primary), Norm(X_aux)) ∈ ℝ^(H×W×(C₁+C₂))', '(3-5)'),

    ('p',
     '主干网络在X_fused上端到端训练，从而显式建模多模态联合表示。'
     '框架最终输出整图像素级分类预测，'
     '评估上既考察整体精度OA、平均精度AA与Kappa系数，'
     '也面向广义零样本学习考察已知类精度（S）、未知类精度（U）'
     '以及二者的调和均值H=2·S·U/(S+U)，'
     '其中调和均值衡量已知类与未知类性能的均衡性，是该任务下最具代表性的综合指标。'),
]


def _set_run_font(run, font_name='宋体', size_pt=12, bold=False):
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def _make_block(doc, item):
    level = item[0]
    if level in ('h1', 'h2', 'h3', 'p', 'fig', 'cap'):
        text = item[1]
    elif level == 'eq':
        text = item[1]
        eq_num = item[2]
    else:
        raise ValueError(level)

    para = doc.add_paragraph()

    if level == 'h1':
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(text)
        _set_run_font(run, font_name='黑体', size_pt=16, bold=True)

    elif level == 'h2':
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(text)
        _set_run_font(run, font_name='黑体', size_pt=15, bold=True)

    elif level == 'h3':
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(text)
        _set_run_font(run, font_name='黑体', size_pt=14, bold=True)

    elif level == 'p':
        para.paragraph_format.first_line_indent = Cm(0.74)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(text)
        _set_run_font(run, font_name='宋体', size_pt=12, bold=False)

    elif level == 'fig':
        # 图占位：居中
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(text)
        _set_run_font(run, font_name='宋体', size_pt=12, bold=False)

    elif level == 'cap':
        # 图题：居中、五号
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text)
        _set_run_font(run, font_name='宋体', size_pt=10.5, bold=False)

    elif level == 'eq':
        # 公式：居中，公式编号右侧用制表/空格
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_before = Pt(2)
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(f'{text}                                  {eq_num}')
        _set_run_font(run, font_name='Times New Roman', size_pt=12, bold=False)
        # 中文括号等仍用宋体作为东亚字体
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    return para._element


def main():
    doc = Document(str(SRC_DOCX))

    # 找到 anchor：第一个以 ANCHOR_TEXT_PREFIX 开头的段落
    anchor = None
    for para in doc.paragraphs:
        if para.text.startswith(ANCHOR_TEXT_PREFIX):
            anchor = para._element
            break
    if anchor is None:
        raise SystemExit(f'未找到 anchor 段落: {ANCHOR_TEXT_PREFIX!r}')

    # 在文档末尾构造新段落，再 detach 后插到 anchor 之前
    new_elements = [_make_block(doc, item) for item in SECTIONS]
    for elem in new_elements:
        elem.getparent().remove(elem)
        anchor.addprevious(elem)

    doc.save(str(DST_DOCX))
    print(f'Inserted {len(SECTIONS)} blocks before anchor "{ANCHOR_TEXT_PREFIX}..." into {DST_DOCX}')


if __name__ == '__main__':
    main()
