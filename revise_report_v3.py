# -*- coding: utf-8 -*-
"""V3: 清理残留的旧版本分析段落。"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document(r'D:\Master_medium\中期实验报告_修订版.docx')


def replace_para_by_text(doc, search, new_text):
    """查找包含搜索文本的段落并替换。"""
    for para in doc.paragraphs:
        if search in para.text:
            for run in para.runs:
                run.text = ''
            if para.runs:
                run = para.runs[0]
                run.text = new_text
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
            else:
                run = para.add_run(new_text)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.size = Pt(12)
            return True
    return False


# 用更独特的片段精确定位旧段落
FIXES = [
    # OLD MUUFL detailed table paragraph (表3分析)
    {
        'search': '但在水上的精度为0%，导致该组合的总体精度OA降至81.63%',
        'new': 'MUUFL数据集上，"树木"（第1类）与"水"（第6类）组合中，UC在树木上取得85.91%的精度，但水体精度为0%。水体在遥感图像中呈现为深色均匀区域，其光谱特征与建筑阴影、道路暗区等地物存在混淆可能。MUUFL的辅助模态为LiDAR，水体在近红外波段的强吸收特性导致LiDAR回波信号极弱或缺失，这一物理现象与CLIP基于光学图像学到的水体视觉表征（深色、反光、波纹等）差异显著，CLIP难以在LiDAR模态中建立有效的水体语义映射。"以草地为主"（第2类）与"黄色路缘"（第10类）组合中，UC在黄路缘上精度为0%，H(aa)=53.05%。黄色路缘是道路附属设施，在整幅图像中仅占极少数像素，ERS超像素分割也难以将其独立成块，CLIP输入的图像区域中往往混杂了道路、人行道等背景信息，导致类别特征被稀释。"裸土和沙地"（第4类）与"布制覆盖板"（第11类）组合中，布制覆盖板的精度同样为0%。布制覆盖板（cloth panels）属于遥感场景中极为罕见的地物，CLIP的训练数据中可能不存在与之对应的有效图像-文本对。在(1,6)组合下，UC在公路上的精度为82.49%，在建筑上为86.84%，而ExViT在建筑上达98.01%，MFT在公路上达97.07%，已知类上的差距仍然明显。'
    },
    # OLD Berlin detailed table paragraph (表4分析)
    {
        'search': 'UC的总体精度OA达到91.59%，为该数据集所有组合中最高',
        'new': 'Berlin数据集上，"土壤"（第5类）与"家庭菜园"（第6类）组合取得了该数据集最高的OA（91.59%），UC在土壤上精度为66.72%，但家庭菜园精度为0%。"土壤"（第5类）与"水"（第8类）组合中，OA=92.47%，土壤精度为66.54%，水同样为0%。水体和家庭菜园在SAR图像中均缺乏独特且稳定的散射特征——水体镜面反射导致的低回波与其他低散射区域（如平整裸地）容易混淆，而家庭菜园的体散射特征与森林、低矮植物等植被类别高度重叠。两者的共同问题是CLIP缺乏对SAR成像机制的物理感知，仅凭转换后的RGB图像难以建立可靠的语义对应。"森林"（第1类）与"家庭菜园"（第6类）组合中，UC的OA降至86.30%，森林精度仅36.93%，家庭菜园仍为0%。(5,6)组合下UC在居民区（98.13%）和工业区（91.43%）上的已知类精度均超过90%，但两个未知类的精度分别为66.72%和0%，这一巨大反差说明UC方法在已知类的特征学习上并无障碍，其性能差异完全来源于CLIP伪标签在特定类别上的质量缺陷。'
    },
    # OLD conclusion paragraph (综合上述分析...第一...第二...第三...)
    {
        'search': '综合上述分析可以得出以下认识：第一，CLIP对遥感场景中面积较小或特征不显著的地物',
        'new': '从上述逐类分析可以看出，CLIP伪标签的质量受三类因素制约较为明显。一是地物的空间尺度：面积较小或分散分布的地物（如黄色路缘仅占少量像素），超像素分割难以将其作为独立区域提取，CLIP输入图像中混杂了大量背景信息，导致类别信号被淹没。二是辅助模态的类型：SAR图像基于微波散射机制成像，与CLIP所学光学图像的视觉统计特性存在本质差异，Berlin数据集多个类别精度为零的现象即源于此；LiDAR虽然也存在模态差异，但其强度图像在视觉上更接近灰度光学图像，与CLIP的兼容性相对更好。三是类别命名的语义覆盖：allotment（家庭菜园）和cloth panels（布制覆盖板）这类在通用图像-文本数据中极少出现的词汇，从根本上就无法触发CLIP的有效语义匹配，表现为相关类别的零精度结果。'
    },
    # OLD beta paragraph
    {
        'search': 'β值越大，模型对伪标签的关注度与依赖程度越高',
        'new': '最后，图8展示了不同伪标签损失权重β下的分类精度变化。β控制着模型在多大程度上信任伪标签提供的监督信号。当β=0时，模型仅使用已知类监督标签训练，退化到纯监督模式，未知类精度几乎为零。β取0.25或0.5时，模型在多数数据集上取得了S(AA)与U(AA)之间的较好平衡。β过大（如1.0或2.0）时，模型对噪声伪标签的过度信任开始损害已知类的分类精度，这在Houston和Berlin数据集上表现尤为明显——S(AA)出现下降的同时U(AA)并未继续提升。MUUFL数据集对β的变化相对不敏感，这可能与其初始伪标签质量整体较高、伪标签噪声水平较低有关。综合来看，β=0.5在多数场景下提供了较好的已知类与未知类性能折中。'
    },
]


def main():
    print('V3: 清理残留旧段落...')
    cleaned = 0
    for i, fix in enumerate(FIXES):
        if replace_para_by_text(doc, fix['search'], fix['new']):
            cleaned += 1
            print(f'  [OK] #{i+1}: {fix["search"][:60]}...')
        else:
            print(f'  [WARN] #{i+1} not found: {fix["search"][:60]}...')

    output_path = r'D:\Master_medium\中期实验报告_修订版.docx'
    doc.save(output_path)
    print(f'\nV3完成！清理 {cleaned}/{len(FIXES)} 处残留段落。')


if __name__ == '__main__':
    main()
