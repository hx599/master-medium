# -*- coding: utf-8 -*-
"""把开题报告 §3.1（主要研究内容）、§3.2.1（基于 Transformer 的多模态遥感图像预训练方法研究）、
§4（预期达到的目标）三节迁移到 中期实验报告_最终版.docx。

执行步骤：
1. 替换 §3 下面那条简短的占位过渡句 → §3.1（含完整三段说明 + 三个子方向的描述）
2. 替换 §3.2 下面那条简短的占位过渡句 → §3.2 真正的引入段（含图 3-1 占位 + caption）
3. 在 §3.2 引入段后插入 §3.2.1 完整内容（含图 3-2 占位 + caption）
4. 在 §3.2.2 最后一段（"主干网络在X_fused上端到端训练..."）后、"表1显示..."之前，插入 §4
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


SRC_DOCX = Path('/Users/agiuser/master-medium/中期实验报告_最终版.docx')
DST_DOCX = SRC_DOCX

# ── 锚点段落（用于精确定位） ─────────────────────────────
ANCHOR_S3_PLACEHOLDER = '针对多模态遥感图像在实际分类应用中面临的对标注样本依赖性强'
ANCHOR_S32_PLACEHOLDER = '本课题的研究方案分为三个子方向'
ANCHOR_AFTER_S322 = '主干网络在X_fused上端到端训练'
ANCHOR_TABLE1 = '表1显示了两个未知类别情况下的各方法平均分类结果'


# ── §3.1 内容 ──────────────────────────────────────
S31_BLOCKS = [
    ('h2', '3.1 主要研究内容'),
    ('p',
     '针对多模态遥感图像在实际分类应用中面临的对标注样本依赖性强、模态间特征差异显著以及在实际应用下难以有效识别未知类别等挑战，'
     '本课题将深入研究视觉语言模型辅助下多模态遥感图像像素级分类，旨在提升模型在多模态数据条件下的跨模态融合能力与未知类别识别能力，'
     '并增强在动态环境中的持续适应性。研究内容分为三个方面：'
     '第一部分主要研究基于Transformer的多模态遥感图像预训练方法，'
     '为多模态数据融合和跨模态特征建模提供坚实的特征基础，为后续未知类分类打下基础；'
     '第二部分主要研究基于视觉语言模型的多模态遥感图像未知类分类方法，'
     '利用视觉语言模型的先验知识解决多模态遥感图像中未见类别的分类问题；'
     '第三部分在第二部分的基础上，针对遥感数据规模庞大、类别持续增加的实际情况，'
     '主要研究基于持续学习的多模态遥感图像未知类分类方法，'
     '以确保在新的未知类别不断出现的情况下，模型保持分类性能的稳定性，并实现对新的未知类别的分类。'),
    ('p',
     '（1）基于Transformer的多模态遥感图像预训练方法。'
     '针对传统多模态遥感图像分类方法（如CNN方法）在建模模态间的长距离依赖关系方面存在不足，'
     '本课题首先研究基于Transformer模型的多模态遥感图像特征融合方法，'
     '设计适用于不同模态（如高光谱、LiDAR、SAR、光学等）特征提取的跨模态Transformer结构。'
     '为了解决Transformer缺少局部归纳偏置导致分类精度低的问题，本课题将研究基于跨模态掩蔽重建任务的预训练方法，'
     '通过在不同模态数据间进行联合重建，使得Transformer模型学习到模态间的互补特征关系，'
     '从而同时获得局部建模能力和跨模态关联建模能力。'),
    ('p',
     '（2）基于视觉语言模型的多模态遥感图像未知类分类方法。'
     '现有多模态遥感图像分类方法普遍假设测试数据中所有类别均在训练阶段出现，'
     '因为无法有效识别训练阶段中未出现的新类别。'
     '针对这一问题，本课题将研究把对象级分类的视觉语言模型迁移至像素级多模态遥感图像分类的技术路线，'
     '探索利用预训练视觉语言模型的先验知识，'
     '在多模态特征空间中建立已知类别、未知类别与图像样本特征之间的映射关系，'
     '从而在保持已知类别高精度分类的同时，实现未知类别的有效分类。'),
    ('p',
     '（3）基于持续学习的多模态遥感图像未知类分类方法。'
     '现有多模态遥感图像分类方法在面对持续到来的新任务和新类别时，往往假设训练阶段能够覆盖所有目标类别，'
     '这使得模型在实际动态环境中难以适应不断扩展的类别集合。'
     '同时，传统方法在学习新类别的过程中容易产生灾难性遗忘，导致对已学类别的分类精度显著下降，'
     '从而限制了模型在真实遥感应用中的长期有效性。针对上述问题，'
     '本课题将研究基于持续学习的多模态遥感图像未知类分类方法。'
     '具体而言，将在已有的视觉语言模型多模态分类框架基础上，'
     '引入混合专家适配器实现模型动态扩展，并结合激活-冻结平衡策略，'
     '在持续任务中保持旧知识的判别能力，同时高效学习新未知类别特征，对分阶段到来的未知类进行持续分类。'
     '该方法旨在构建能够在任务序列中不断自适应的分类模型，实现已知类的稳健分类与未知类的有效持续学习。'),
]


# ── §3.2 引入段（替换占位）──────────────────────────────
S32_INTRO_BLOCKS = [
    ('fig', '[图 3-1 占位]'),
    ('cap', '图3-1　总体研究方案'),
    ('p',
     '本课题的总体研究方案如图3-1所示，包括基于Transformer的多模态遥感图像预训练方法、'
     '基于视觉语言模型的多模态遥感图像未知类分类方法和基于持续学习的多模态遥感图像未知类分类方法三部分。'
     '第一部分，采用掩蔽图像重建策略对Transformer进行预训练，'
     '为后续未知类分类提供特征提取能力强的Transformer模型；'
     '第二部分利用视觉语言模型解决多模态遥感图像的未知类分类问题，'
     '为基于持续学习的未知类分类提供研究基础；'
     '最后对多模态遥感图像持续学习分类问题进行研究。'),
]


# ── §3.2.1 完整内容 ─────────────────────────────────
S321_BLOCKS = [
    ('h3', '3.2.1 基于Transformer的多模态遥感图像预训练方法研究'),
    ('p',
     '在多模态遥感图像特征提取方法中，卷积神经网络采用局部连接和权值共享的方式，能够有效进行特征提取，'
     '在早期研究中占据主导地位。然而，基于CNN的多模态遥感图像分类方法由于卷积核感受野受限，'
     '难以对不同模态数据中长程依赖关系进行有效建模。'
     '针对跨模态和跨场景条件下复杂的空间分布和数据差异，需要同时考虑局部和全局的像素依赖关系，'
     '建立跨模态像素乃至地物之间的空间上下文关系。'
     'Transformer模型的出现为长距离依赖建模与全局特征提取提供了强有力的工具。'
     '因此，本课题将以Transformer为基础，实现多模态遥感图像的局部和长程特征联合提取，'
     '为后续未知类分类与持续学习任务奠定特征提取基础。'),
    ('p',
     '多模态遥感数据不仅在空间和光谱特征上差异显著，还具有不同的噪声特性和数据分辨率差异。'
     '传统视觉Transformer在输入时大多将图像划分为空间块作为序列输入，'
     '主要依赖空间自注意力机制，而忽视了不同模态特征的互补性和特有属性。'
     '因此，直接应用这些模型难以在多模态遥感特征提取中获得理想的实验精度。'),
    ('fig', '[图 3-2 占位]'),
    ('cap', '图3-2　基于Transformer的多模态遥感图像预训练流程示意图'),
    ('p',
     '为了充分利用多模态遥感图像中的空间-光谱信息，本课题拟在空间块划分之前，'
     '先对光谱信息进行处理，设计一个光谱信息聚合模块，对不同模态数据的光谱特征进行初步提取，'
     '确保每个空间块中蕴含丰富的光谱信息，使后续Transformer模型能够高效提取融合特征。'
     '与此同时，遥感图像具有局部空间一致性，即相近空间位置的像素更可能属于同一地物类别。'
     '本课题还将研究基于掩蔽重建的多模态遥感图像预训练方法，通过空间掩蔽任务，'
     '使Transformer从掩蔽的图像中恢复原始信息，'
     '使得Transformer在预训练阶段提升对局部空间信息的关注，提升模型学习局部与模态的独立特征。'),
    ('p',
     '然而，单纯依靠掩蔽重建获得的特征虽然具备一定判别能力，但主要关注输入样本的语义信息，'
     '其对于不同训练样本的区分性不足，难以直接满足分类需求。'
     '因此，本课题将引入实例对比学习与类间判别监督两种机制，'
     '将判别性特征提取能力注入Transformer的预训练阶段，'
     '获得既具有全局建模能力又具有强判别性的多模态Transformer特征提取模型，'
     '整体预训练流程如图3-2所示。其中实例对比学习用于区分不同样本，'
     '类间判别监督用于强调不同类别间的差异性。'
     '该模型将在后续未知类分类与持续学习阶段通过微调的方法进行应用。'),
]


# ── §4 完整内容 ──────────────────────────────────
S4_BLOCKS = [
    ('h1', '4. 预期达到的目标'),
    ('p',
     '（1）通过设计适用于多模态遥感数据的跨模态Transformer结构，'
     '并引入跨模态掩蔽重建与对比学习机制，'
     '探索同时兼具全局建模能力与判别性的特征提取模型，'
     '实现对高光谱、LiDAR、SAR等多模态数据的高效融合和表征。'
     '基于掩蔽重建的预训练方法，构建跨模态视觉Transformer模型来完成多模态遥感数据的分类问题，'
     '为后续未知类分类与持续学习奠定坚实的特征基础。'),
    ('p',
     '（2）通过将预训练视觉语言模型（如CLIP）迁移至像素级分类任务，'
     '并结合超像素伪标签生成与知识蒸馏方法，'
     '建立已知类、未知类与多模态图像特征之间的语义映射关系。'
     '预期实现在缺乏训练样本的情况下对未知类别的有效分类，'
     '同时保证已知类别的高精度分类性能，提升模型的语义泛化与跨模态适应能力。'),
    ('p',
     '（3）基于视觉语言模型的语义先验与混合专家适配器（MoE Adapters）的动态扩展机制，'
     '结合激活-冻结平衡策略，实现对新未知类别的持续学习与对旧知识的长期保持。'
     '预期能够构建一个持续学习框架，'
     '使模型在动态环境中不断适应新任务和新类别的同时，保持对已知类别的识别精度，'
     '并具备对未知类别的持续学习能力，'
     '从而在真实应用场景中实现更可靠和持久的多模态遥感图像分类性能。'),
]


def _set_run_font(run, font_name='宋体', size_pt=12, bold=False):
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def _make_block(doc, item):
    level = item[0]
    text = item[1]

    para = doc.add_paragraph()

    if level == 'h1':
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(text)
        _set_run_font(run, font_name='黑体', size_pt=16, bold=True)

    elif level == 'h2':
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(text)
        _set_run_font(run, font_name='黑体', size_pt=15, bold=True)

    elif level == 'h3':
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(text)
        _set_run_font(run, font_name='黑体', size_pt=14, bold=True)

    elif level == 'p':
        para.paragraph_format.first_line_indent = Cm(0.74)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(text)
        _set_run_font(run, font_name='宋体', size_pt=12, bold=False)

    elif level == 'fig':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(text)
        _set_run_font(run, font_name='宋体', size_pt=12, bold=False)

    elif level == 'cap':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text)
        _set_run_font(run, font_name='宋体', size_pt=10.5, bold=False)

    return para._element


def _find_anchor(doc, prefix):
    for para in doc.paragraphs:
        if para.text.startswith(prefix):
            return para
    raise SystemExit(f'未找到 anchor: {prefix!r}')


def _replace_paragraph_text(para, new_text):
    """清空一个段落的所有 run，然后写入新文本（保留段落自身格式）。"""
    # 先清空 run
    for run in list(para.runs):
        r = run._element
        r.getparent().remove(r)
    # 增加一个新 run
    new_run = para.add_run(new_text)
    _set_run_font(new_run, font_name='宋体', size_pt=12, bold=False)


def main():
    doc = Document(str(SRC_DOCX))

    # 锚点 1：§3 下方的占位段（针对多模态…）
    s3_anchor = _find_anchor(doc, ANCHOR_S3_PLACEHOLDER)
    # 锚点 2：§3.2 下方的占位段（本课题的研究方案分为…）
    s32_anchor = _find_anchor(doc, ANCHOR_S32_PLACEHOLDER)
    # 锚点 3：§3.2.2 末段（主干网络在X_fused上…）—— §4 插入位置
    s322_last = _find_anchor(doc, ANCHOR_AFTER_S322)

    # === Step 1 ===
    # 把 §3 占位段当作 §3.1 的开篇用（删掉占位段，在它之前插入 §3.1 标题 + §3.1 全部内容）
    # 但 §3 自己也需要一句过渡。简单做法：保留 §3 占位段不动（已经是合理的 §3 总览）；
    # 然后在它之后插 §3.1（包含 h2 标题 + 4 段正文，其中第 1 段会与 §3 占位段重复，所以替换占位段为更短的过渡）
    #
    # 调整策略：§3 占位段已经是"针对多模态遥感图像…旨在提升模型的跨模态融合能力与未知类别识别能力，
    # 并增强其在动态环境中的持续适应性。" —— 把它替换为 §3 真正应有的过渡（删掉与 §3.1 第一段重复的部分），
    # 然后在它之后插入 §3.1 完整内容。
    _replace_paragraph_text(
        s3_anchor,
        '本章针对多模态遥感图像在实际分类应用中面临的对标注样本依赖性强、模态间特征差异显著'
        '以及难以有效识别未知类别等挑战，从主要研究内容和研究方案两个维度展开介绍。'
    )

    # 在 s3_anchor 之后插入 §3.1 全部 block（用 addnext 链式倒序插）
    insert_after = s3_anchor._element
    for item in S31_BLOCKS:
        new_elem = _make_block(doc, item)
        new_elem.getparent().remove(new_elem)
        insert_after.addnext(new_elem)
        insert_after = new_elem

    # === Step 2 ===
    # 把 §3.2 占位段替换为 §3.2 真正的过渡，并紧接其后插入图 3-1 + caption + intro
    _replace_paragraph_text(
        s32_anchor,
        '本节给出 §3.1 中三个研究内容的具体研究方案。'
    )

    insert_after = s32_anchor._element
    for item in S32_INTRO_BLOCKS:
        new_elem = _make_block(doc, item)
        new_elem.getparent().remove(new_elem)
        insert_after.addnext(new_elem)
        insert_after = new_elem

    # 紧接其后插入 §3.2.1 全部 block
    for item in S321_BLOCKS:
        new_elem = _make_block(doc, item)
        new_elem.getparent().remove(new_elem)
        insert_after.addnext(new_elem)
        insert_after = new_elem

    # === Step 3 ===
    # §4 插入到 §3.2.2 最后一段之后
    insert_after = s322_last._element
    for item in S4_BLOCKS:
        new_elem = _make_block(doc, item)
        new_elem.getparent().remove(new_elem)
        insert_after.addnext(new_elem)
        insert_after = new_elem

    doc.save(str(DST_DOCX))
    print(f'§3.1 / §3.2 intro / §3.2.1 / §4 inserted into {DST_DOCX}')


if __name__ == '__main__':
    main()
