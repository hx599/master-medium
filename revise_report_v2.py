# -*- coding: utf-8 -*-
"""
修改中期实验报告中的分析文字，降低AI痕迹 - V2。
改用段落索引直接定位 + 文本匹配结合的方式。
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

doc = Document(r'D:\Master_medium\中期实验报告_修订版.docx')


def replace_para(para, new_text):
    """替换段落文字，保留字体格式。"""
    for run in para.runs:
        run.text = ''
    if para.runs:
        run = para.runs[0]
        run.text = new_text
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)
    else:
        run = para.add_run(new_text)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(12)


def find_para_by_start(doc, start_text):
    """查找以指定文本开头的段落。"""
    for para in doc.paragraphs:
        if para.text.strip().startswith(start_text):
            return para
    return None


def find_para_by_contains(doc, text_fragment):
    """查找包含指定文本片段的段落。"""
    for para in doc.paragraphs:
        if text_fragment in para.text:
            return para
    return None


# ── 使用段落开头匹配的替换 ──
REPLACEMENTS_BY_START = [
    # Houston per-class unseen analysis (图1)
    {
        'start': '在Houston数据集上，UC方法表现最好的组合包括',
        'new': 'Houston数据集上，表现较好的组合包括"受压草地"（第2类）与"高速公路"（第10类）（U(aa)=69.78%，H(aa)=76.62%）、"商业建筑"（第8类）与"铁路"（第11类）（U(aa)=76.43%，H(aa)=76.43%）等。这两组未知类在视觉形态上差异显著——受压草地呈现不均匀的绿色植被纹理，而高速公路呈现宽阔的灰色线性特征；商业建筑为块状结构，铁路则呈现细长线性特征——CLIP对不同视觉概念的语义匹配相对准确。相比之下，"网球场"（第14类）几乎在所有涉及它的组合中均导致模型失效：与健康草地组合时U(aa)=0.71%，与土壤组合时U(aa)=0.67%，与人造草地组合时U(aa)=0.81%。网球场在高分辨率遥感图像中通常呈现为绿色矩形区域，其纹理和色调与草地类地物高度相似，CLIP的通用视觉先验难以在缺乏光谱辅助信息的情况下做出可靠区分。类似问题也出现在公路与铁路的组合中（U(aa)=2.69%，H(aa)=5.21%），两类地物在遥感图像中均呈现为灰色线性结构，光谱和空间特征重叠严重。'
    },
    # MUUFL per-class unseen analysis (图2)
    {
        'start': 'MUUFL数据集上的整体表现优于Houston，这与其类别命名更贴近日常用语有关',
        'new': 'MUUFL数据集上的整体表现优于Houston。树木与泥沙地组合取得了U(aa)=82.78%、H(aa)=74.77%的较好结果，树木与建筑物组合也有U(aa)=71.98%、H(aa)=70.47%。但"混合地表"（第3类）成为该数据集上的主要难点：与水体组合时U(aa)仅9.52%，与人行道组合时U(aa)仅2.10%，与黄色路缘组合时U(aa)降至1.75%。这一类别本身定义模糊，涵盖了多种地表覆盖类型，CLIP难以在语义空间中为其建立一致的映射。"水体"作为未知类时同样表现不佳——与人行道组合U(aa)=6.17%，与黄色路缘组合U(aa)=4.50%。水体在LiDAR数据中呈现独特的低反射特征，这一物理特性在CLIP基于光学图像训练得到的视觉表征中无法得到有效表达，导致了跨模态的信息丢失。'
    },
    # Berlin per-class unseen analysis (图3)
    {
        'start': 'Berlin数据集是三个数据集中最困难的，这与其采用SAR辅助模态以及部分类别命名较为抽象密切相关',
        'new': 'Berlin数据集是三个数据集中最具挑战性的。表现相对最好的组合是"低矮植物"（第4类）与"土壤"（第5类）（U(aa)=37.77%，H(aa)=52.38%），以及"工业区"（第3类）与"土壤"（第5类）（U(aa)=35.28%，H(aa)=49.93%）。土壤类别在SAR图像中的散射特性相对稳定，当与光谱区分度较高的类别配对时伪标签质量稍好，但绝对精度仍然有限。居住区在多数组合中都表现出较低的未知类精度——与低矮植物组合时U(aa)=8.18%，与家庭菜园组合时U(aa)=1.60%，与商业区组合时U(aa)=2.33%。居住区在SAR图像中的散射纹理与工业区、商业区等建成区类别存在较大重叠，呈现为相似的亮色斑点状分布，加之"residential area"在CLIP的语义空间中缺乏高度具象的视觉锚点，伪标签难以达到可用的精度水平。最极端的情况出现在家庭菜园与商业区的组合中（U(aa)=0.25%，H(aa)=0.50%），模型几乎完全丧失了对这两类未知地物的识别能力。"allotment"（家庭菜园）在日常英语中使用频率较低，CLIP训练数据中可能极少存在与之对应的遥感图像-文本对，从根本上限制了伪标签的生成质量。'
    },
]

# ── 使用段落索引 + 内容验证的替换（解决编码匹配问题）──
REPLACEMENTS_BY_INDEX = [
    # Houston detailed table analysis (表2)
    {
        'verify': '被压草',
        'new': 'Houston数据集上，"被压草"（第2类）与"高速路"（第10类）组合取得了H(aa)=76.79%的较好结果，UC在高速路上的精度为66.52%，在被压草上为73.82%。被压草呈现不均匀的绿色植被纹理，高速路呈现宽阔的灰色线性特征，二者的视觉差异为CLIP的语义匹配提供了清晰依据。但"树木"（第4类）与"商业区"（第8类）组合的H(aa)仅43.15%，UC在商业区上的精度为48.99%。该组合还出现了已知类精度受扰动的现象——高速路作为已知类的精度降至8.20%，公路仅为66.88%。这可能由于商业建筑的屋顶光谱与周边地物存在较大重叠，其伪标签中的噪声通过光谱特征修正环节传播到了部分视觉相似的已知类别。在"铁路"（第11类）与"跑道"（第15类）组合中，UC在铁路上的精度为57.20%，H(aa)=70.16%。(11,15)组合的H(aa)高于(4,8)组合，说明铁路与跑道虽然在局部尺度上存在视觉相似性（均为灰色线性结构），但二者的空间上下文差异（铁路通常贯穿多个区域，跑道局限于固定场地）为模型区分提供了一定的空间线索。ExViT和MFT在已知类上均保持了90%以上的精度，在(4,8)组合下ExViT对居民区的精度达97.49%，MFT对网球场的精度达99.53%，表明传统方法的能力瓶颈集中在零样本分类环节。'
    },
    # MUUFL detailed table analysis (表3)
    {
        'verify': 'MUUFL',
        'new': 'MUUFL数据集上，"树木"（第1类）与"水"（第6类）组合中，UC在树木上取得85.91%的精度，但水体精度为0%。水体在遥感图像中呈现为深色均匀区域，其光谱特征与建筑阴影、道路暗区等地物存在混淆可能。MUUFL的辅助模态为LiDAR，水体在近红外波段的强吸收特性导致LiDAR回波信号极弱或缺失，这一物理现象与CLIP基于光学图像学到的水体视觉表征（深色、反光、波纹等）差异显著，CLIP难以在LiDAR模态中建立有效的水体语义映射。"以草地为主"（第2类）与"黄色路缘"（第10类）组合中，UC在黄路缘上精度为0%，H(aa)=53.05%。黄色路缘是道路附属设施，在整幅图像中仅占极少数像素，ERS超像素分割也难以将其独立成块，CLIP输入的图像区域中往往混杂了道路、人行道等背景信息，导致类别特征被稀释。"裸土和沙地"（第4类）与"布制覆盖板"（第11类）组合中，布制覆盖板的精度同样为0%。布制覆盖板（cloth panels）属于遥感场景中极为罕见的地物，CLIP的训练数据中可能不存在与之对应的有效图像-文本对。在(1,6)组合下，UC在公路上的精度为82.49%，在建筑上为86.84%，而ExViT在建筑上达98.01%，MFT在公路上达97.07%，已知类上的差距仍然明显。'
    },
    # Berlin detailed table analysis (表4)
    {
        'verify': '家庭菜园',
        'new': 'Berlin数据集上，"土壤"（第5类）与"家庭菜园"（第6类）组合取得了该数据集最高的OA（91.59%），UC在土壤上精度为66.72%，但家庭菜园精度为0%。"土壤"（第5类）与"水"（第8类）组合中，OA=92.47%，土壤精度为66.54%，水同样为0%。水体和家庭菜园在SAR图像中均缺乏独特且稳定的散射特征——水体镜面反射导致的低回波与其他低散射区域（如平整裸地）容易混淆，而家庭菜园的体散射特征与森林、低矮植物等植被类别高度重叠。两者的共同问题是CLIP缺乏对SAR成像机制的物理感知，仅凭转换后的RGB图像难以建立可靠的语义对应。"森林"（第1类）与"家庭菜园"（第6类）组合中，UC的OA降至86.30%，森林精度仅36.93%，家庭菜园仍为0%。(5,6)组合下UC在居民区（98.13%）和工业区（91.43%）上的已知类精度均超过90%，但两个未知类的精度分别为66.72%和0%，这一巨大反差说明UC方法在已知类的特征学习上并无障碍，其性能差异完全来源于CLIP伪标签在特定类别上的质量缺陷。'
    },
    # Beta parameter analysis (图8)
    {
        'verify': '损失权重',
        'new': '最后，图8展示了不同伪标签损失权重β下的分类精度变化。β控制着模型在多大程度上信任伪标签提供的监督信号。当β=0时，模型仅使用已知类监督标签训练，退化到纯监督模式，未知类精度几乎为零。β取0.25或0.5时，模型在多数数据集上取得了S(AA)与U(AA)之间的较好平衡。β过大（如1.0或2.0）时，模型对噪声伪标签的过度信任开始损害已知类的分类精度，这在Houston和Berlin数据集上表现尤为明显——S(AA)出现下降的同时U(AA)并未继续提升。MUUFL数据集对β的变化相对不敏感，这可能与其初始伪标签质量整体较高、伪标签噪声水平较低有关。综合来看，β=0.5在多数场景下提供了较好的已知类与未知类性能折中。'
    },
]


def main():
    print('正在修改中期实验报告 (V2)...')
    modified = 0

    # Phase 1: start-text matching
    for i, rule in enumerate(REPLACEMENTS_BY_START):
        para = find_para_by_start(doc, rule['start'])
        if para is not None:
            replace_para(para, rule['new'])
            modified += 1
            print(f'  [OK] start-match #{i+1}')
        else:
            print(f'  [WARN] start-match #{i+1} not found')

    # Phase 2: content fragment matching (for paragraphs where startswith failed)
    for i, rule in enumerate(REPLACEMENTS_BY_INDEX):
        para = find_para_by_contains(doc, rule['verify'])
        if para is not None:
            replace_para(para, rule['new'])
            modified += 1
            print(f'  [OK] content-match #{i+1}: verify="{rule["verify"]}"')
        else:
            print(f'  [WARN] content-match #{i+1}: verify="{rule["verify"]}" not found')

    output_path = r'D:\Master_medium\中期实验报告_修订版.docx'
    doc.save(output_path)
    print(f'\n修改完成！共替换 {modified}/{len(REPLACEMENTS_BY_START) + len(REPLACEMENTS_BY_INDEX)} 处分析段落。')
    print(f'修订版已保存至: {output_path}')


if __name__ == '__main__':
    main()
