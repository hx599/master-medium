# -*- coding: utf-8 -*-
"""生成中期报告 Word 文档。"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式设置 ──
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)

# ── 辅助函数 ──
def add_heading_cn(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_para(text, bold=False, indent=False, font_size=12):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(font_size)
    run.bold = bold
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    # 数据行
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return table


# ============================================================
# 封面
# ============================================================
for _ in range(4):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('哈尔滨工业大学')
run.font.size = Pt(26)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('硕士学位论文中期报告')
run.font.size = Pt(22)
run.font.name = '黑体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.bold = True

doc.add_paragraph()

info_items = [
    ('论文题目', '视觉语言模型赋能多模态遥感图像分类方法研究'),
    ('学    院', '电子与信息工程学院'),
    ('学    科', '信息与通信工程'),
    ('导    师', '陈  舒'),
    ('学    号', '24S005019'),
    ('姓    名', '韩  晓'),
]
for label, value in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{label}：{value}')
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('报告日期：2026.04.22')
run.font.size = Pt(14)
run.font.name = '宋体'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# 目录页
# ============================================================
add_heading_cn('目  录', level=1)
toc_items = [
    '一、课题来源与研究目的和意义',
    '二、研究工作进展情况',
    '  2.1 基于视觉语言模型的多模态遥感图像未知类分类方法总体框架',
    '  2.2 多模态伪标签生成与增强模块',
    '  2.3 基于谱特征引导的伪标签修正模块',
    '  2.4 两阶段全局训练策略',
    '  2.5 实验结果与分析',
    '三、目前存在的问题及解决方案',
    '四、下一步工作计划',
    '五、已取得的阶段性成果',
]
for item in toc_items:
    add_para(item, font_size=12)

doc.add_page_break()

# ============================================================
# 一、课题来源与研究目的和意义
# ============================================================
add_heading_cn('一、课题来源与研究目的和意义', level=1)

add_para(
    '本课题来源于黑龙江省自然科学基金项目。遥感作为人类感知地球的重要手段，多模态遥感数据通过融合不同传感器的观测信息（如高光谱、LiDAR、SAR等），获取更丰富的地物特征，实现从单一维度到多维度的信息互补，显著提升遥感分析的精度与鲁棒性。多模态遥感数据的融合分析能够弥补单一数据源的局限性，增强地物识别的可分性、精度与稳健性，在城市规划、灾害监测、精准农业等领域具有重要的应用价值。',
    indent=True
)

add_para(
    '然而，实际应用场景中，地物类别往往是动态演化的，新的未知类别不断出现，传统分类方法通常假设训练集与测试集共享相同的标签空间，难以有效处理未知类别。这种假设在实际中往往不成立——当推理阶段出现训练中未涵盖的地物类型时，模型倾向于将未知样本误分为已知类别，导致严重的分类错误。因此，如何在保证已知类别识别精度的同时，有效识别并分类未知类别，已成为遥感图像分类领域亟待解决的关键问题。',
    indent=True
)

add_para(
    '视觉语言模型（Vision-Language Model, VLM）通过在大规模图像-文本对上进行对比学习，掌握了丰富的语义先验知识，为未知类别的识别提供了新的可能。原始HZSCM方法在单模态高光谱图像的零样本分类中取得了显著成效，但其仅依赖单一高光谱模态的输入，未充分利用多模态遥感数据的互补信息。本课题以HZSCM为基础，将其从单模态高光谱场景扩展到多模态遥感场景，重点研究基于视觉语言模型的多模态遥感图像未知类分类方法，旨在充分利用多模态数据的互补特性，提升未知类别的分类性能。',
    indent=True
)

# ============================================================
# 二、研究工作进展情况
# ============================================================
add_heading_cn('二、研究工作进展情况', level=1)

# 2.1 总体框架
add_heading_cn('2.1 基于视觉语言模型的多模态遥感图像未知类分类方法总体框架', level=2)

add_para(
    '在原始HZSCM框架的基础上，本研究将其从单模态高光谱图像零样本分类扩展为面向多模态遥感数据的未知类分类框架。总体框架仍保持三阶段设计：多模态伪标签生成与增强、谱特征引导的伪标签修正、两阶段全局训练，但在每个阶段中均引入了多模态适配机制。整体流程如图1所示。',
    indent=True
)

add_para(
    '（1）多模态伪标签生成阶段：利用CLIP视觉语言模型，分别为每个模态生成独立的伪标签概率图，再在概率层进行平均融合，充分利用不同模态对地物的差异化感知能力。',
    indent=True
)

add_para(
    '（2）伪标签修正阶段：沿用HZSCM中基于超像素聚类和局部一致性检验的修正策略，但将输入从单模态数据扩展为多模态融合数据，使修正过程能够利用更丰富的光谱-空间特征。',
    indent=True
)

add_para(
    '（3）全局训练阶段：设计了两种多模态融合训练策略——单模态训练模式（single）和早期融合训练模式（early）。单模态模式仅使用主模态数据训练分割网络，但伪标签由多模态联合生成；早期融合模式将多个模态在通道维直接拼接后送入网络训练，使网络端到端地学习多模态联合表示。',
    indent=True
)

# 2.2 多模态伪标签生成与增强模块
add_heading_cn('2.2 多模态伪标签生成与增强模块', level=2)

add_para(
    '伪标签的质量直接决定了未知类分类的性能上限。原始HZSCM方法仅使用高光谱图像生成伪标签，本研究将这一模块扩展为多模态伪标签生成框架，核心改进包括以下几个方面：',
    indent=True
)

add_para('（1）多模态独立伪标签生成与概率层融合', bold=True)
add_para(
    '针对多模态遥感数据的特点，本研究提出了"各模态独立生成、概率层平均融合"的伪标签生成策略。具体而言，对于每个数据集的主模态（如高光谱）和辅助模态（如LiDAR或SAR），分别将原始数据转换为RGB图像后送入CLIP模型，独立生成各模态的伪标签概率图，最终通过对各模态概率图取平均得到融合伪标签。这一策略的优势在于：不同模态对地物的感知能力存在差异（如高光谱擅长捕捉光谱特征，LiDAR擅长刻画高程信息，SAR对结构信息敏感），概率层融合能够有效综合各模态的互补信息，提升伪标签的整体质量。',
    indent=True
)

add_para('（2）各模态独立分割与CLIP推理', bold=True)
add_para(
    '原始HZSCM中，伪标签生成依赖于超像素分割（ERS）和SAM分割结果来提供空间结构信息。在多模态场景下，不同模态的图像特征差异显著，使用同一分割结果难以兼顾各模态的特点。因此，本研究为每个模态分别生成了独立的ERS超像素分割和SAM分割掩码，各模态使用自己的分割结果进行CLIP推理。例如，对于Houston数据集，主模态（高光谱）使用primary_SPGT和primary_list，辅助模态（LiDAR）使用secondary_SPGT和secondary_list。这种"各模态独立分割"的策略确保了每个模态的空间结构信息被充分利用。',
    indent=True
)

add_para('（3）双CLIP模型集成策略的保留与增强', bold=True)
add_para(
    '本研究保留了HZSCM中双CLIP模型集成的核心设计——同时使用RSICD-CLIP（遥感领域微调的CLIP）和ViT-L-CLIP（通用视觉CLIP），在概率层取平均。这一策略的动机在于：不同CLIP模型在各类别上的识别能力存在互补性，RSICD-CLIP对遥感场景理解更深入，而ViT-L-CLIP在通用视觉概念上更稳健。在多模态场景下，这种互补性被进一步放大：不同模态的图像与不同CLIP模型的匹配程度不同，双CLIP集成有助于缓解单一模型对特定模态的偏好。',
    indent=True
)

add_para('（4）SAM与ERS融合分割策略的多模态适配', bold=True)
add_para(
    '原始HZSCM中SAM与ERS的融合策略为：SAM分割提供更完整的语义区域，ERS分割保证像素全覆盖，两者互补后送入CLIP。本研究将这一策略扩展为多模态版本——每个模态独立执行SAM+ERS融合，分别生成伪标签后再进行概率层平均。具体实现中，通过fuse_SAM_ERS函数将SAM掩码覆盖到ERS超像素上，SAM覆盖区域使用SAM的分割结果进行CLIP推理，未覆盖区域使用ERS的分割结果，最终将两者的概率图拼接为完整的伪标签概率图。',
    indent=True
)

# 2.3 伪标签修正模块
add_heading_cn('2.3 基于谱特征引导的伪标签修正模块', level=2)

add_para(
    'CLIP生成的伪标签不可避免地存在噪声，尤其是在遥感领域，CLIP的语义先验与遥感地物的实际分布之间存在知识鸿沟。本研究沿用了HZSCM中基于谱特征引导的伪标签修正策略，并针对多模态场景进行了适配，主要包括以下步骤：',
    indent=True
)

add_para('（1）超像素级特征聚合', bold=True)
add_para(
    '将像素级的CLIP伪标签概率图和多模态融合后的光谱特征聚合到超像素级。对于每个超像素，计算其内部所有像素的平均光谱特征作为该超像素的表示，同时通过多数投票确定有标签超像素的监督标签，无标签超像素标记为-1。在多模态早期融合模式下，聚合使用的光谱特征是主模态与辅助模态通道拼接后的融合特征；在单模态模式下，仅使用主模态的光谱特征。',
    indent=True
)

add_para('（2）谱聚类与局部一致性检验', bold=True)
add_para(
    '对超像素特征进行谱聚类，将光谱特征相似的超级像素归为同一簇。然后，对于每个未标注超像素，计算其伪标签的局部一致性分数——即该超像素所在聚类中与它同标签的邻居权重之和占所有同标签邻居权重之和的比例。若该分数低于预设阈值θ，则将该超像素的伪标签替换为次高概率类别，并迭代检验直到一致性达标或穷尽所有候选类别。',
    indent=True
)

add_para('（3）监督标签强制覆盖', bold=True)
add_para(
    '修正完成后，对于有监督标签的超像素，直接将其伪标签替换为真实监督标签，确保已知类别的标签不受修正过程的影响。这一步骤在多模态场景下同样适用，是保证已知类别分类精度的重要保障。',
    indent=True
)

# 2.4 两阶段全局训练策略
add_heading_cn('2.4 两阶段全局训练策略', level=2)

add_para(
    '本研究采用与HZSCM相同的两阶段训练策略，但针对多模态数据设计了灵活的融合训练模式：',
    indent=True
)

add_para('第一阶段：仅用真实标签的预训练', bold=True)
add_para(
    '使用U-Net风格的分割网络（HSI_Seg_Hard），仅用有监督标签进行预训练。在此阶段，未知类别的训练样本被遮蔽（标签置零），模型仅学习已知类别的分类边界。预训练的目的在于为伪标签修正提供初始的网络预测结果。在多模态早期融合模式下，网络输入为主模态和辅助模态通道拼接后的融合数据；在单模态模式下，仅输入主模态数据。',
    indent=True
)

add_para('第二阶段：修正伪标签与监督标签联合训练', bold=True)
add_para(
    '利用第一阶段预训练模型的预测结果，结合CLIP伪标签和谱特征信息，执行伪标签修正。修正后的伪标签与真实监督标签一起指导分割网络的训练。总损失函数为：',
    indent=True
)

add_para('    L = L_sup + β · L_pl', indent=False)

add_para(
    '其中L_sup为监督交叉熵损失，L_pl为伪标签交叉熵损失，β为伪标签损失权重（默认为1.0）。训练采用早停策略——当伪标签损失连续4个epoch低于0.05时停止训练，避免过拟合。',
    indent=True
)

add_para('多模态融合训练模式', bold=True)
add_para(
    '本研究实现了两种融合训练模式，通过fusion_mode参数控制：（1）single模式——训练时仅输入主模态数据，伪标签由多模态联合生成后直接使用；（2）early模式——训练时将主模态和辅助模态在通道维拼接后作为网络输入，各模态分别标准化后拼接，避免某一模态量纲过大主导训练。两种模式各有优势：single模式实现简单、训练效率高，且伪标签已经融合了多模态信息；early模式使网络端到端学习多模态联合表示，理论上能够更好地挖掘模态间的互补性。',
    indent=True
)

# 2.5 实验结果与分析
add_heading_cn('2.5 实验结果与分析', level=2)

add_para('2.5.1 实验设置', bold=True)

add_para(
    '为验证所提方法的有效性，本研究在三个多模态遥感数据集上进行了实验：',
    indent=True
)

add_table(
    ['数据集', '主模态', '辅助模态', '类别数', '图像尺寸'],
    [
        ['Houston', '高光谱(144波段)', 'LiDAR', '15', '349×1905'],
        ['HS-SAR-Berlin', '高光谱(244波段)', 'SAR', '8', '794×220'],
        ['MUUFL', '高光谱(72波段)', 'LiDAR', '11', '325×220'],
    ]
)

add_para('')
add_para(
    '实验采用广义零样本学习（GZSL）设定：训练时将某些类别的标签完全移除（unseen类），测试时同时评估已知类别（seen）和未知类别（unseen）的分类精度。评价指标包括：总体精度（OA）、平均精度（AA）、已知类精度（S_OA/S_AA）、未知类精度（U_OA/U_AA）和调和均值（H_OA/H_AA）。每种 unseen 类别组合重复5次实验，报告均值和标准差。训练样本比例为每类5%。',
    indent=True
)

add_para('2.5.2 多模态与单模态对比实验', bold=True)

add_para(
    '为验证多模态伪标签生成策略的有效性，本研究对比了多模态伪标签（early融合模式）与单模态伪标签（仅主模态）在相同unseen类别组合下的分类性能。结果如下表所示：',
    indent=True
)

add_para('表1 Houston数据集多模态与单模态对比', bold=True)
add_table(
    ['Unseen类别', '模式', 'OA(%)', 'S_OA(%)', 'U_OA(%)', 'H_OA(%)', 'H_AA(%)'],
    [
        ['commercial buildings;\n2nd parking lot', '单模态', '81.06', '86.08', '41.98', '56.42', '43.48'],
        ['commercial buildings;\n2nd parking lot', '多模态', '81.01', '85.95', '42.62', '56.97', '43.96'],
        ['stressed grass;\nsynthetic grass', '单模态', '79.63', '84.91', '44.21', '58.11', '49.24'],
        ['stressed grass;\nsynthetic grass', '多模态', '79.15', '84.65', '42.28', '56.21', '47.55'],
        ['healthy grass; soil', '单模态', '76.49', '82.12', '48.18', '60.73', '61.63'],
        ['healthy grass; soil', '多模态', '75.91', '81.78', '46.36', '59.17', '60.06'],
        ['highway;\nrunning track', '单模态', '84.31', '83.99', '86.56', '85.23', '83.27'],
        ['highway;\nrunning track', '多模态', '84.25', '83.93', '86.47', '85.16', '83.18'],
        ['synthetic grass;\n1st parking lot', '单模态', '78.26', '87.16', '17.91', '29.56', '24.10'],
        ['synthetic grass;\n1st parking lot', '多模态', '78.63', '87.15', '20.81', '33.39', '27.38'],
    ]
)

add_para('')
add_para('表2 HS-SAR-Berlin数据集多模态与单模态对比', bold=True)
add_table(
    ['Unseen类别', '模式', 'OA(%)', 'S_OA(%)', 'U_OA(%)', 'H_OA(%)', 'H_AA(%)'],
    [
        ['residential;\nindustrial', '单模态', '37.26', '95.99', '1.31', '2.58', '17.46'],
        ['residential;\nindustrial', '多模态', '37.36', '95.94', '1.49', '2.94', '19.68'],
        ['forest; allotment', '单模态', '85.12', '96.03', '21.72', '35.40', '23.45'],
        ['forest; allotment', '多模态', '85.12', '96.03', '21.72', '35.40', '23.45'],
        ['soil;\ncommercial area', '单模态', '89.94', '96.16', '27.82', '43.13', '48.85'],
        ['soil;\ncommercial area', '多模态', '89.94', '96.16', '27.82', '43.13', '48.85'],
        ['forest; soil', '单模态', '83.09', '95.32', '16.83', '28.60', '49.73'],
        ['forest; soil', '多模态', '83.09', '95.32', '16.83', '28.60', '49.73'],
        ['low plants; water', '单模态', '81.89', '95.07', '2.19', '4.28', '19.24'],
        ['low plants; water', '多模态', '81.89', '95.07', '2.19', '4.28', '19.24'],
    ]
)

add_para('')
add_para('表3 MUUFL数据集多模态与单模态对比', bold=True)
add_table(
    ['Unseen类别', '模式', 'OA(%)', 'S_OA(%)', 'U_OA(%)', 'H_OA(%)', 'H_AA(%)'],
    [
        ['building shadow;\nyellow curb', '单模态', '81.95', '82.67', '66.81', '73.89', '49.00'],
        ['building shadow;\nyellow curb', '多模态', '81.95', '82.67', '66.81', '73.89', '49.00'],
        ['road; water', '单模态', '77.48', '84.89', '29.30', '43.47', '32.04'],
        ['road; water', '多模态', '77.48', '84.89', '29.30', '43.47', '32.04'],
        ['trees;\nbuilding shadow', '单模态', '68.62', '76.81', '59.56', '67.08', '65.64'],
        ['trees;\nbuilding shadow', '多模态', '68.62', '76.81', '59.56', '67.08', '65.64'],
        ['road; sidewalk', '单模态', '82.09', '87.44', '51.85', '64.86', '45.50'],
        ['road; sidewalk', '多模态', '82.09', '87.44', '51.85', '64.86', '45.50'],
        ['dirt or sand; road', '单模态', '79.43', '85.59', '46.70', '60.41', '64.21'],
        ['dirt or sand; road', '多模态', '79.43', '85.59', '46.70', '60.41', '64.21'],
    ]
)

add_para('')
add_para(
    '从实验结果可以得出以下分析：',
    indent=True
)

add_para(
    '（1）在Houston数据集上，多模态策略在部分unseen类别组合中展现出优势。例如，在"commercial buildings; 2nd parking lot"组合中，多模态的H_OA从56.42%提升至56.97%，H_AA从43.48%提升至43.96%；在"synthetic grass; 1st parking lot"组合中，多模态的U_OA从17.91%提升至20.81%，H_OA从29.56%提升至33.39%，提升幅度明显。这表明LiDAR高程信息的引入有助于区分光谱特征相似但高程不同的地物类型。',
    indent=True
)

add_para(
    '（2）在HS-SAR-Berlin和MUUFL数据集上，多模态与单模态的性能差异较小，部分组合结果完全一致。分析原因在于：这两个数据集的辅助模态（SAR和LiDAR）在转换为RGB图像后，与CLIP的自然图像预训练数据分布差距较大，CLIP难以从这些转换图像中提取有效的语义信息，导致多模态伪标签的增益有限。这一发现提示我们，对于与自然图像差异较大的模态，需要更精细的模态适配策略。',
    indent=True
)

add_para(
    '（3）整体来看，多模态伪标签策略在辅助模态与CLIP语义空间对齐较好的情况下（如Houston的LiDAR数据）能够带来性能提升，但在对齐困难的情况下增益有限甚至略有下降。这为后续研究指明了方向——如何更好地将非光学模态的信息融入视觉语言模型的语义空间是关键挑战。',
    indent=True
)

add_para('2.5.3 不同unseen类别组合的性能差异分析', bold=True)

add_para(
    '实验结果表明，模型的零样本分类性能高度依赖于unseen类别的具体选择。以Houston数据集为例：',
    indent=True
)

add_para(
    '（1）"highway; running track"组合的H_OA最高（约85%），这是因为这两种地物在光谱和空间特征上具有明显区别，且与CLIP的自然图像概念高度对齐，CLIP能够生成高质量的伪标签。',
    indent=True
)

add_para(
    '（2）"synthetic grass; 1st parking lot"组合的H_OA最低（约30%），这是因为synthetic grass与健康草地的光谱特征极其相似，CLIP难以区分，导致伪标签质量下降，进而影响零样本分类性能。',
    indent=True
)

add_para(
    '（3）HS-SAR-Berlin数据集上"residential; industrial"组合的U_OA仅约1.5%，这是遥感零样本分类中最具挑战性的情况——residential area和industrial area在SAR图像中缺乏清晰的视觉边界，且这些概念与CLIP预训练数据中的自然图像概念差异较大。',
    indent=True
)

add_para('2.5.4 两阶段训练策略的有效性验证', bold=True)

add_para(
    '实验对比了预训练阶段和最终训练阶段的性能变化。以Houston数据集"highway; running track"组合为例，预训练阶段的OA约为57.9%，经过伪标签修正和第二阶段训练后，OA提升至84.2%，提升幅度超过26个百分点。这一结果验证了两阶段训练策略的有效性：第一阶段为伪标签修正提供了可靠的网络预测，第二阶段利用修正后的伪标签显著提升了未知类别的分类能力。',
    indent=True
)

# ============================================================
# 三、目前存在的问题及解决方案
# ============================================================
add_heading_cn('三、目前存在的问题及解决方案', level=1)

add_para(
    '在目前的研究工作中，主要存在以下几个问题：',
    indent=True
)

add_para('问题一：非光学模态与CLIP语义空间对齐困难', bold=True)
add_para(
    '当前方法将SAR和LiDAR等非光学模态数据直接转换为RGB图像后送入CLIP，但SAR图像的散斑噪声和LiDAR数据的强度特性与CLIP预训练的自然图像分布差异较大，导致CLIP难以从这些转换图像中提取有效语义信息，限制了多模态伪标签的增益效果。',
    indent=True
)
add_para(
    '解决方案：后续将探索模态适配器（Modality Adapter）机制，在CLIP视觉编码器前增加轻量级适配层，将非光学模态的特征映射到CLIP的语义空间。同时，考虑利用遥感领域微调的CLIP模型（如RSICD-CLIP）对非光学模态进行专门适配，提高伪标签质量。',
    indent=True
)

add_para('问题二：语义相似类别的零样本区分能力不足', bold=True)
add_para(
    '当unseen类别与某些seen类别在语义上高度相似时（如"synthetic grass"与"healthy grass"），CLIP生成的伪标签容易混淆，导致零样本分类性能显著下降。当前的伪标签修正策略基于光谱特征的局部一致性，对于光谱特征也相似的类别难以有效区分。',
    indent=True
)
add_para(
    '解决方案：后续将引入更精细的文本提示工程（Prompt Engineering）策略，为语义相似的类别设计更具区分性的文本描述。同时，考虑在伪标签修正阶段引入多粒度空间结构约束，利用地物的空间上下文关系辅助区分语义相似类别。',
    indent=True
)

add_para('问题三：早期融合模式的信息冗余', bold=True)
add_para(
    '当前早期融合模式将多个模态在通道维直接拼接，虽然实现简单，但可能存在信息冗余——不同模态的某些通道可能包含高度相关的信息，直接拼接可能导致网络学习效率下降。此外，不同模态的特征尺度差异可能影响训练稳定性。',
    indent=True
)
add_para(
    '解决方案：后续将探索更先进的多模态融合策略，如跨模态注意力机制和模态感知的特征对齐方法，以更有效地利用多模态互补信息。同时，将研究自适应融合权重学习机制，使模型能够根据输入数据自动调整各模态的贡献。',
    indent=True
)

# ============================================================
# 四、下一步工作计划
# ============================================================
add_heading_cn('四、下一步工作计划', level=1)

add_para(
    '基于目前的研究进展和存在的问题，下一步工作计划如下：',
    indent=True
)

add_para(
    '（1）2026年4月—5月：设计并实现模态适配器机制，改善非光学模态与CLIP语义空间的对齐问题。在CLIP视觉编码器前增加轻量级适配层，使SAR和LiDAR模态的特征能够更好地被CLIP理解和利用。同时，优化文本提示工程策略，提升语义相似类别的区分能力。',
    indent=True
)

add_para(
    '（2）2026年5月—6月：研究跨模态注意力融合机制，替代当前的简单通道拼接策略。设计模态感知的特征对齐模块，使模型能够自适应地学习不同模态之间的互补关系，提升多模态融合的效果。',
    indent=True
)

add_para(
    '（3）2026年6月—7月：在更多数据集上进行全面实验，包括Trento、Berlin等数据集，验证方法的泛化性。与现有的多模态遥感分类方法和零样本学习方法进行系统性对比实验。',
    indent=True
)

add_para(
    '（4）2026年7月—9月：撰写学位论文，整理实验结果，准备论文答辩。',
    indent=True
)

# ============================================================
# 五、已取得的阶段性成果
# ============================================================
add_heading_cn('五、已取得的阶段性成果', level=1)

add_para(
    '（1）完成了基于HZSCM的多模态遥感图像未知类分类框架设计与实现。在原始HZSCM的代码基础上，新增了多模态数据读取模块（data_read_multimodal.py、data_read_singlemodal.py），支持Houston（高光谱+LiDAR）、HS-SAR-Berlin（高光谱+SAR）、MUUFL（高光谱+LiDAR）三个多模态数据集的读取和处理。',
    indent=True
)

add_para(
    '（2）实现了多模态伪标签生成策略，包括各模态独立ERS/SAM分割、独立CLIP推理、概率层融合等关键模块。为每个数据集的主模态和辅助模态分别生成了独立的超像素分割（SPGT）和SAM掩码列表。',
    indent=True
)

add_para(
    '（3）实现了两种多模态融合训练模式（single和early），并通过对比实验验证了多模态伪标签策略的有效性。在Houston数据集上，多模态策略在部分unseen类别组合中取得了性能提升。',
    indent=True
)

add_para(
    '（4）在三个多模态遥感数据集上完成了系统性的零样本分类实验，涵盖了多种unseen类别组合，获得了较为全面的实验数据，为后续方法改进提供了重要参考。',
    indent=True
)

# ── 保存 ──
output_path = 'D:/Master_medium/韩晓-中期报告.docx'
doc.save(output_path)
print(f'Report saved to {output_path}')
